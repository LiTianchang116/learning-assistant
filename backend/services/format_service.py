import json
import os
import uuid
import zipfile

from docx import Document

from config import TEMP_DIR
from services.llm_client import chat
from tools.docx_tools import TOOL_SCHEMAS, execute_tool
from memory import get_context_messages, add_message

SYSTEM_PROMPT = """你是一个Word文档格式修改助手。你可以使用工具来修改文档的字体、字号、加粗、颜色、对齐、标题、页面边距、行间距、字间距等。

当用户要求修改"所有段落"或"所有正文"时，先用 list_paragraphs 查看文档结构，再对所有段落索引执行操作。
当用户提到"小四号"时，对应12磅；"五号"对应10.5磅；"四号"对应14磅。
修改前先查询文档结构，修改后用中文总结做了什么。"""


def get_doc_summary(doc: Document) -> str:
    lines = [f"段落数: {len(doc.paragraphs)}", "段落列表:"]
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else "Normal"
        text = p.text[:50].replace("\n", " ")
        lines.append(f"  [{i}] {style} | {text}")
    return "\n".join(lines)


def process_format_request(file_paths: list[str], user_message: str, session_id: str) -> dict:
    results = {}
    for fpath in file_paths:
        doc = Document(fpath)
        summary = get_doc_summary(doc)
        context = get_context_messages(session_id, "format")
        messages = [{"role": "system", "content": SYSTEM_PROMPT + "\n\n当前文档:\n" + summary}]
        messages.extend(context)
        messages.append({"role": "user", "content": user_message})

        add_message(session_id, "format", "user", user_message)

        for _ in range(10):
            resp = chat(messages, tools=TOOL_SCHEMAS)
            content = resp.get("content", "")
            tool_calls = resp.get("tool_calls", [])

            if not tool_calls:
                add_message(session_id, "format", "assistant", content)
                results[fpath] = {"summary": content, "doc": doc}
                break

            # Build assistant message with content blocks (Anthropic format)
            assistant_content = []
            if content:
                assistant_content.append({"type": "text", "text": content})
            for tc in tool_calls:
                raw_args = tc["function"]["arguments"]
                func_args = raw_args if isinstance(raw_args, dict) else json.loads(raw_args)
                assistant_content.append({
                    "type": "tool_use",
                    "id": tc["id"],
                    "name": tc["function"]["name"],
                    "input": func_args,
                })
            messages.append({"role": "assistant", "content": assistant_content})

            # Build tool results as user message (Anthropic format)
            tool_results = []
            for tc in tool_calls:
                func_name = tc["function"]["name"]
                raw_args = tc["function"]["arguments"]
                func_args = raw_args if isinstance(raw_args, dict) else json.loads(raw_args)
                result = execute_tool(doc, func_name, func_args)
                tool_results.append({
                    "type": "tool_result",
                    "tool_use_id": tc["id"],
                    "content": result,
                })
            messages.append({"role": "user", "content": tool_results})
        else:
            add_message(session_id, "format", "assistant", "已达到最大操作次数")
            results[fpath] = {"summary": "已达到最大操作次数", "doc": doc}

    # Save modified files
    output_paths = []
    for fpath, info in results.items():
        out_name = f"modified_{uuid.uuid4().hex[:8]}_{os.path.basename(fpath)}"
        out_path = os.path.join(TEMP_DIR, out_name)
        info["doc"].save(out_path)
        output_paths.append(out_path)

    # If multiple files, zip them
    if len(output_paths) > 1:
        zip_name = f"modified_{uuid.uuid4().hex[:8]}.zip"
        zip_path = os.path.join(TEMP_DIR, zip_name)
        with zipfile.ZipFile(zip_path, "w") as zf:
            for p in output_paths:
                zf.write(p, os.path.basename(p))
        return {"download_path": zip_path, "filename": zip_name, "summaries": {os.path.basename(k): v["summary"] for k, v in results.items()}}

    return {
        "download_path": output_paths[0],
        "filename": os.path.basename(output_paths[0]),
        "summaries": {os.path.basename(k): v["summary"] for k, v in results.items()},
    }
