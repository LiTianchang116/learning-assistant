import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

def _make_tool(name: str, desc: str, props: dict, required: list) -> dict:
    return {"name": name, "description": desc, "input_schema": {"type": "object", "properties": props, "required": required}}


TOOL_SCHEMAS = [
    _make_tool("set_font", "设置文档中指定段落的字体名称",
        {"paragraph_indices": {"type": "array", "items": {"type": "integer"}, "description": "段落索引列表"},
         "font_name": {"type": "string", "description": "字体名称，如 '宋体', 'Times New Roman'"}},
        ["paragraph_indices", "font_name"]),
    _make_tool("set_font_size", "设置文档中指定段落的字号",
        {"paragraph_indices": {"type": "array", "items": {"type": "integer"}, "description": "段落索引列表"},
         "size_pt": {"type": "number", "description": "字号（磅），如12为小四号"}},
        ["paragraph_indices", "size_pt"]),
    _make_tool("set_bold", "设置或取消加粗",
        {"paragraph_indices": {"type": "array", "items": {"type": "integer"}, "description": "段落索引列表"},
         "bold": {"type": "boolean", "description": "true加粗，false取消"}},
        ["paragraph_indices", "bold"]),
    _make_tool("set_color", "设置字体颜色",
        {"paragraph_indices": {"type": "array", "items": {"type": "integer"}, "description": "段落索引列表"},
         "color_hex": {"type": "string", "description": "颜色十六进制，如 #FF0000"}},
        ["paragraph_indices", "color_hex"]),
    _make_tool("set_alignment", "设置段落对齐方式",
        {"paragraph_indices": {"type": "array", "items": {"type": "integer"}, "description": "段落索引列表"},
         "alignment": {"type": "string", "enum": ["left", "center", "right", "justify"], "description": "对齐方式"}},
        ["paragraph_indices", "alignment"]),
    _make_tool("set_heading", "设置段落为标题样式",
        {"paragraph_indices": {"type": "array", "items": {"type": "integer"}, "description": "段落索引列表"},
         "level": {"type": "integer", "description": "标题级别1-9，0为正文"}},
        ["paragraph_indices", "level"]),
    _make_tool("set_page_margins", "设置页面边距",
        {"top_cm": {"type": "number", "description": "上边距（厘米）"},
         "bottom_cm": {"type": "number", "description": "下边距（厘米）"},
         "left_cm": {"type": "number", "description": "左边距（厘米）"},
         "right_cm": {"type": "number", "description": "右边距（厘米）"}},
        []),
    _make_tool("set_line_spacing", "设置行间距",
        {"paragraph_indices": {"type": "array", "items": {"type": "integer"}, "description": "段落索引列表"},
         "spacing_rule": {"type": "string", "enum": ["multiple", "exactly", "at_least"], "description": "multiple=倍数, exactly=固定值(磅)"},
         "spacing_value": {"type": "number", "description": "倍数值或磅值"}},
        ["paragraph_indices", "spacing_rule", "spacing_value"]),
    _make_tool("set_character_spacing", "设置字间距",
        {"paragraph_indices": {"type": "array", "items": {"type": "integer"}, "description": "段落索引列表"},
         "spacing_pt": {"type": "number", "description": "字间距（磅）"}},
        ["paragraph_indices", "spacing_pt"]),
    _make_tool("list_paragraphs", "列出文档所有段落的索引、样式和文本预览",
        {"style_filter": {"type": "string", "description": "按样式名筛选，如 'Heading'"}},
        []),
]


def execute_tool(doc: Document, name: str, args: dict) -> str:
    if name == "list_paragraphs":
        return _list_paragraphs(doc, args.get("style_filter"))
    indices = args.get("paragraph_indices", [])
    total = len(doc.paragraphs)
    for i in indices:
        if i < 0 or i >= total:
            return f"Error: 段落索引 {i} 超出范围（共 {total} 段）"

    if name == "set_font":
        return _set_font(doc, indices, args["font_name"])
    elif name == "set_font_size":
        return _set_font_size(doc, indices, args["size_pt"])
    elif name == "set_bold":
        return _set_bold(doc, indices, args["bold"])
    elif name == "set_color":
        return _set_color(doc, indices, args["color_hex"])
    elif name == "set_alignment":
        return _set_alignment(doc, indices, args["alignment"])
    elif name == "set_heading":
        return _set_heading(doc, indices, args["level"])
    elif name == "set_page_margins":
        return _set_page_margins(doc, args)
    elif name == "set_line_spacing":
        return _set_line_spacing(doc, indices, args["spacing_rule"], args["spacing_value"])
    elif name == "set_character_spacing":
        return _set_character_spacing(doc, indices, args["spacing_pt"])
    return f"Error: 未知工具 {name}"


def _list_paragraphs(doc, style_filter=None):
    result = []
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else "Normal"
        if style_filter and style_filter.lower() not in style.lower():
            continue
        text = para.text[:60]
        result.append({"index": i, "style": style, "text": text})
    return json.dumps(result, ensure_ascii=False)


def _apply_to_runs(doc, indices, func):
    for idx in indices:
        para = doc.paragraphs[idx]
        for run in para.runs:
            func(run)
        if not para.runs and para.text:
            func(para.add_run(para.text))


def _set_font(doc, indices, font_name):
    def f(run):
        run.font.name = font_name
        # For CJK fonts
        from docx.oxml.ns import qn
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    _apply_to_runs(doc, indices, f)
    return f"已设置字体为 {font_name}，段落 {indices}"


def _set_font_size(doc, indices, size_pt):
    _apply_to_runs(doc, indices, lambda r: setattr(r.font.size, '_Pt', size_pt) or setattr(r.font, 'size', Pt(size_pt)))
    return f"已设置字号为 {size_pt}pt，段落 {indices}"


def _set_bold(doc, indices, bold):
    _apply_to_runs(doc, indices, lambda r: setattr(r, 'bold', bold))
    return f"已设置加粗={bold}，段落 {indices}"


def _set_color(doc, indices, color_hex):
    h = color_hex.lstrip("#")
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    _apply_to_runs(doc, indices, lambda run: setattr(run.font.color, 'rgb', RGBColor(r, g, b)))
    return f"已设置颜色={color_hex}，段落 {indices}"


def _set_alignment(doc, indices, alignment):
    for idx in indices:
        doc.paragraphs[idx].alignment = ALIGN_MAP.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
    return f"已设置对齐={alignment}，段落 {indices}"


def _set_heading(doc, indices, level):
    for idx in indices:
        if level == 0:
            doc.paragraphs[idx].style = doc.styles["Normal"]
        else:
            doc.paragraphs[idx].style = doc.styles[f"Heading {level}"]
    return f"已设置标题级别={level}，段落 {indices}"


def _set_page_margins(doc, args):
    sec = doc.sections[0]
    if "top_cm" in args:
        sec.top_margin = Cm(args["top_cm"])
    if "bottom_cm" in args:
        sec.bottom_margin = Cm(args["bottom_cm"])
    if "left_cm" in args:
        sec.left_margin = Cm(args["left_cm"])
    if "right_cm" in args:
        sec.right_margin = Cm(args["right_cm"])
    return f"已设置页面边距: {args}"


def _set_line_spacing(doc, indices, rule, value):
    from docx.enum.text import WD_LINE_SPACING
    rule_map = {"multiple": WD_LINE_SPACING.MULTIPLE, "exactly": WD_LINE_SPACING.EXACTLY, "at_least": WD_LINE_SPACING.AT_LEAST}
    for idx in indices:
        pf = doc.paragraphs[idx].paragraph_format
        if rule == "multiple":
            pf.line_spacing = value
        else:
            pf.line_spacing = Pt(value)
        pf.line_spacing_rule = rule_map[rule]
    return f"已设置行间距: {rule}={value}，段落 {indices}"


def _set_character_spacing(doc, indices, spacing_pt):
    from lxml import etree
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    twips = int(spacing_pt * 20)
    for idx in indices:
        para = doc.paragraphs[idx]
        for run in para.runs:
            rpr = run._element.find(f".//{{{ns}}}rPr")
            if rpr is None:
                rpr = etree.SubElement(run._element, f"{{{ns}}}rPr")
            sp = rpr.find(f"{{{ns}}}spacing")
            if sp is None:
                sp = etree.SubElement(rpr, f"{{{ns}}}spacing")
            sp.set(f"{{{ns}}}val", str(twips))
    return f"已设置字间距={spacing_pt}pt，段落 {indices}"
